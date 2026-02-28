# agents/processors/docx_processor.py
"""
Extracts text and embedded images from DOCX files using python-docx.
Also handles legacy .doc files with a graceful fallback message.
"""

import io
from dataclasses import dataclass, field
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


MAX_TEXT_CHARS = 12_000

# Minimum image dimensions to bother extracting
MIN_IMAGE_WIDTH  = 80
MIN_IMAGE_HEIGHT = 80


@dataclass
class ExtractedImage:
    bytes: bytes
    mime_type: str
    width: int
    height: int
    index: int          # sequential index within the document


@dataclass
class DOCXExtractionResult:
    text: str
    summary: str
    was_truncated: bool
    images: list[ExtractedImage] = field(default_factory=list)
    error: Optional[str] = None


def extract_docx(file_bytes: bytes, mime_type: str) -> DOCXExtractionResult:
    """
    Extract text and images from DOCX bytes.
    Legacy .doc files cannot be processed — returns helpful error.
    """
    # .doc (legacy binary) format is not supported by python-docx
    if mime_type == "application/msword":
        return DOCXExtractionResult(
            text="",
            summary="",
            was_truncated=False,
            error="Legacy .doc format is not supported. Please convert to .docx and re-upload."
        )

    if not DOCX_AVAILABLE:
        return DOCXExtractionResult(
            text="",
            summary="",
            was_truncated=False,
            error="python-docx not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        return DOCXExtractionResult(
            text="",
            summary="",
            was_truncated=False,
            error=f"Could not open DOCX: {e}"
        )

    text_parts: list[str] = []
    total_chars = 0
    was_truncated = False

    # ── text extraction ───────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if was_truncated:
            break

        text = para.text.strip()
        if not text:
            continue

        # Detect headings by style name
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            chunk = f"\n\n## {text}\n"
        else:
            chunk = f"\n{text}"

        if total_chars + len(chunk) > MAX_TEXT_CHARS:
            remaining = MAX_TEXT_CHARS - total_chars
            if remaining > 50:
                text_parts.append(chunk[:remaining])
            text_parts.append("\n\n[... content truncated. Document continues beyond this point.]")
            was_truncated = True
            total_chars = MAX_TEXT_CHARS
        else:
            text_parts.append(chunk)
            total_chars += len(chunk)

    # Also extract table text
    if not was_truncated:
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    chunk = " | ".join(row_cells) + "\n"
                    if total_chars + len(chunk) > MAX_TEXT_CHARS:
                        was_truncated = True
                        break
                    text_parts.append(chunk)
                    total_chars += len(chunk)
            if was_truncated:
                break

    # ── image extraction ──────────────────────────────────────────────────────
    images: list[ExtractedImage] = []
    img_index = 0

    try:
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                img_part = rel.target_part
                img_bytes = img_part.blob
                content_type = img_part.content_type  # e.g. "image/jpeg"

                if content_type not in {"image/jpeg", "image/png", "image/gif", "image/webp"}:
                    continue

                # Attempt to get dimensions via PIL if available
                width, height = _get_image_dimensions(img_bytes)
                if width < MIN_IMAGE_WIDTH or height < MIN_IMAGE_HEIGHT:
                    continue

                images.append(ExtractedImage(
                    bytes=img_bytes,
                    mime_type=content_type,
                    width=width,
                    height=height,
                    index=img_index,
                ))
                img_index += 1
            except Exception:
                continue
    except Exception:
        pass  # if image extraction fails entirely, still return text

    full_text = "".join(text_parts).strip()
    summary   = full_text[:800] if full_text else ""

    return DOCXExtractionResult(
        text=full_text,
        summary=summary,
        was_truncated=was_truncated,
        images=images,
    )


def _get_image_dimensions(img_bytes: bytes) -> tuple[int, int]:
    """Returns (width, height). Falls back to (999, 999) if PIL unavailable."""
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(img_bytes))
        return img.size  # (width, height)
    except Exception:
        return (999, 999)  # assume large enough if we can't check